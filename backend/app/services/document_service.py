from typing import List, Optional, Dict, Any, Tuple
import tiktoken
import uuid
import logging
from pypdf import PdfReader
from docx import Document as DocxDocument
import io

from app.models.document import Document, DocumentChunk
from app.services.pinecone_service import get_pinecone_service

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for processing and managing documents."""

    def __init__(self):
        self.pinecone = get_pinecone_service()
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.max_chunk_tokens = 500
        self.chunk_overlap = 50
        # In-memory cache for recently uploaded documents
        self._documents: Dict[str, Document] = {}

    # =========================================================================
    # Private helpers for Pinecone operations
    # =========================================================================

    async def _fetch_chunks_from_pinecone(self, document_id: str) -> List[Dict[str, Any]]:
        """
        Fetch and sort chunks for a document from Pinecone.

        Args:
            document_id: The document ID to fetch chunks for

        Returns:
            List of chunks sorted by chunk_index, or empty list if not found
        """
        try:
            chunks = await self.pinecone.get_chunks_by_document(document_id)
            if not chunks:
                return []

            # Sort by chunk_index for consistent ordering
            return sorted(
                chunks,
                key=lambda c: c.get("metadata", {}).get("chunk_index", 0)
            )
        except Exception as e:
            logger.error(f"Failed to fetch chunks from Pinecone for {document_id}: {e}")
            return []

    def _reconstruct_document_from_chunks(
        self,
        document_id: str,
        chunks: List[Dict[str, Any]]
    ) -> Optional[Document]:
        """
        Reconstruct a Document object from Pinecone chunks.

        Args:
            document_id: The document ID
            chunks: List of chunks (should already be sorted)

        Returns:
            Document object or None if chunks are empty
        """
        if not chunks:
            return None

        first_chunk = chunks[0]
        metadata = first_chunk.get("metadata", {})

        return Document(
            id=document_id,
            filename=metadata.get("filename", "Unknown Document"),
            content_type=metadata.get("content_type", "application/pdf"),
            file_size=metadata.get("file_size", 0),
            chunk_count=len(chunks),
            chunk_ids=[chunk["id"] for chunk in chunks]
        )

    def _count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        return len(self.tokenizer.encode(text))

    def _chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Split text into chunks while preserving structure.

        Returns list of dicts with 'content' and 'type' (heading, paragraph, list).
        """
        # Split into paragraphs (preserve structure)
        paragraphs = self._split_into_paragraphs(text)

        chunks = []
        current_chunk_parts = []
        current_tokens = 0

        for para in paragraphs:
            para_text = para["content"]
            para_tokens = self._count_tokens(para_text)

            # If single paragraph exceeds limit, split by sentences
            if para_tokens > self.max_chunk_tokens:
                # Flush current chunk first
                if current_chunk_parts:
                    chunks.append(self._merge_chunk_parts(current_chunk_parts))
                    current_chunk_parts = []
                    current_tokens = 0

                # Split large paragraph into sentence-based chunks
                sentence_chunks = self._split_paragraph_by_sentences(para)
                chunks.extend(sentence_chunks)
            elif current_tokens + para_tokens > self.max_chunk_tokens:
                # Flush current chunk and start new one
                if current_chunk_parts:
                    chunks.append(self._merge_chunk_parts(current_chunk_parts))
                current_chunk_parts = [para]
                current_tokens = para_tokens
            else:
                # Add to current chunk
                current_chunk_parts.append(para)
                current_tokens += para_tokens

        # Flush remaining
        if current_chunk_parts:
            chunks.append(self._merge_chunk_parts(current_chunk_parts))

        return chunks

    def _split_into_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """Split text into paragraphs with type detection."""
        paragraphs = []
        # Split on double newlines or single newlines followed by patterns
        raw_paragraphs = text.split('\n\n')

        for raw in raw_paragraphs:
            raw = raw.strip()
            if not raw:
                continue

            # Further split on single newlines that indicate structure
            lines = raw.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                para_type = self._detect_paragraph_type(line)
                paragraphs.append({
                    "content": line,
                    "type": para_type
                })

        return paragraphs

    def _detect_paragraph_type(self, text: str) -> str:
        """Detect the type of a text block."""
        text_stripped = text.strip()

        # Heading patterns (short, often title case or all caps)
        if len(text_stripped) < 100:
            # All caps heading
            if text_stripped.isupper() and len(text_stripped.split()) <= 10:
                return "heading"
            # Numbered heading (e.g., "1. Introduction", "Chapter 2")
            if text_stripped[:2].replace('.', '').isdigit():
                return "heading"
            # Title case and short
            words = text_stripped.split()
            if len(words) <= 8 and sum(1 for w in words if w[0].isupper()) >= len(words) * 0.6:
                return "heading"

        # List item patterns
        if text_stripped.startswith(('•', '-', '*', '●', '○')):
            return "list_item"
        if len(text_stripped) > 2 and text_stripped[0].isdigit() and text_stripped[1] in '.):':
            return "list_item"

        return "paragraph"

    def _split_paragraph_by_sentences(self, para: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split a large paragraph into sentence-based chunks."""
        text = para["content"]
        para_type = para["type"]

        # Simple sentence split (handles common cases)
        sentences = []
        current = ""
        for char in text:
            current += char
            if char in '.!?' and len(current) > 20:
                sentences.append(current.strip())
                current = ""
        if current.strip():
            sentences.append(current.strip())

        # Group sentences into chunks
        chunks = []
        current_chunk = []
        current_tokens = 0

        for sentence in sentences:
            sent_tokens = self._count_tokens(sentence)
            if current_tokens + sent_tokens > self.max_chunk_tokens:
                if current_chunk:
                    chunks.append({
                        "content": ' '.join(current_chunk),
                        "type": para_type
                    })
                current_chunk = [sentence]
                current_tokens = sent_tokens
            else:
                current_chunk.append(sentence)
                current_tokens += sent_tokens

        if current_chunk:
            chunks.append({
                "content": ' '.join(current_chunk),
                "type": para_type
            })

        return chunks

    def _merge_chunk_parts(self, parts: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Merge multiple paragraph parts into a single chunk."""
        if not parts:
            return {"content": "", "type": "paragraph"}

        # Join with double newlines to preserve structure
        content = '\n\n'.join(p["content"] for p in parts)

        # Determine dominant type
        types = [p["type"] for p in parts]
        if types[0] == "heading":
            chunk_type = "heading"
        elif "list_item" in types and types.count("list_item") > len(types) / 2:
            chunk_type = "list"
        else:
            chunk_type = "paragraph"

        return {"content": content, "type": chunk_type}

    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content."""
        reader = PdfReader(io.BytesIO(content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        doc = DocxDocument(io.BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def _extract_text(self, content: bytes, content_type: str, filename: str) -> str:
        """Extract text from document based on type."""
        if content_type == "application/pdf" or filename.endswith(".pdf"):
            return self._extract_text_from_pdf(content)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.endswith(".docx"):
            return self._extract_text_from_docx(content)
        elif content_type.startswith("text/") or filename.endswith((".txt", ".md")):
            return content.decode("utf-8")
        else:
            raise ValueError(f"Unsupported content type: {content_type}")

    async def process_document(
        self,
        filename: str,
        content: bytes,
        content_type: str
    ) -> Document:
        """
        Process and embed a document.

        Args:
            filename: Name of the file
            content: Raw file content
            content_type: MIME type of the file

        Returns:
            Document object with metadata
        """
        # Create document record
        doc_id = str(uuid.uuid4())
        document = Document(
            id=doc_id,
            filename=filename,
            content_type=content_type,
            file_size=len(content)
        )

        # Extract text
        text = self._extract_text(content, content_type, filename)

        # Chunk the text (now returns structured chunks with type)
        chunks = self._chunk_text(text)
        document.chunk_count = len(chunks)

        # Create embeddings for chunk content
        chunk_contents = [c["content"] for c in chunks]
        embeddings = await self.pinecone.create_embeddings_batch(chunk_contents)

        # Prepare vectors for Pinecone with structure metadata
        vectors = []
        chunk_ids = []
        for i, (chunk_data, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_id = f"{doc_id}_{i}"
            chunk_ids.append(chunk_id)
            vectors.append({
                "id": chunk_id,
                "values": embedding,
                "metadata": {
                    "document_id": doc_id,
                    "filename": filename,
                    "chunk_index": i,
                    "content": chunk_data["content"],
                    "content_type": chunk_data["type"]  # heading, paragraph, list
                }
            })

        # Upsert to Pinecone
        await self.pinecone.upsert_vectors(vectors)

        # Update document record
        document.chunk_ids = chunk_ids
        document.is_embedded = True

        # Store document
        self._documents[doc_id] = document

        return document

    async def delete_document(self, document_id: str) -> bool:
        """
        Delete a document and its vectors.

        Args:
            document_id: ID of the document to delete

        Returns:
            True if deleted, False if not found
        """
        # Try cache first, then Pinecone
        document = self._documents.get(document_id)
        if not document:
            chunks = await self._fetch_chunks_from_pinecone(document_id)
            document = self._reconstruct_document_from_chunks(document_id, chunks)

        if not document:
            return False

        # Delete vectors from Pinecone
        await self.pinecone.delete_vectors(ids=document.chunk_ids)

        # Remove from cache if present
        self._documents.pop(document_id, None)
        return True

    async def get_document(self, document_id: str) -> Optional[Document]:
        """
        Get a document by ID.

        Checks in-memory cache first, falls back to Pinecone reconstruction.
        """
        # Check cache first
        if document := self._documents.get(document_id):
            return document

        # Reconstruct from Pinecone
        chunks = await self._fetch_chunks_from_pinecone(document_id)
        return self._reconstruct_document_from_chunks(document_id, chunks)

    def list_documents(self) -> List[Document]:
        """List all documents."""
        return list(self._documents.values())

    async def search_documents(
        self,
        query: str,
        top_k: int = 5,
        document_ids: Optional[List[str]] = None
    ) -> str:
        """
        Search documents and return formatted context.

        Args:
            query: Search query
            top_k: Number of chunks to retrieve
            document_ids: Optional filter by document IDs

        Returns:
            Formatted context string for the AI
        """
        results = await self.pinecone.search_documents(
            query=query,
            top_k=top_k,
            document_ids=document_ids
        )

        if not results:
            return ""

        context_parts = []
        for result in results:
            if result.get("metadata"):
                filename = result["metadata"].get("filename", "Unknown")
                content = result["metadata"].get("content", "")
                context_parts.append(f"[From {filename}]:\n{content}")

        return "\n\n---\n\n".join(context_parts)

    async def get_document_content(self, document_id: str) -> Dict[str, Any]:
        """
        Get full document content for preview.

        Args:
            document_id: ID of document

        Returns:
            Dictionary with document metadata and full content

        Raises:
            ValueError: If document not found
        """
        # Fetch chunks once - used for both document reconstruction and content
        chunks = await self._fetch_chunks_from_pinecone(document_id)
        if not chunks:
            raise ValueError(f"Document not found: {document_id}")

        # Get or reconstruct document
        document = self._documents.get(document_id)
        if not document:
            document = self._reconstruct_document_from_chunks(document_id, chunks)

        # Build content from chunks (already sorted by _fetch_chunks_from_pinecone)
        chunk_contents = []
        content_parts = []

        for chunk in chunks:
            metadata = chunk.get("metadata", {})
            content = metadata.get("content", "")
            if content:
                chunk_contents.append({
                    "id": chunk["id"],
                    "content": content,
                    "chunk_index": metadata.get("chunk_index", 0),
                    "content_type": metadata.get("content_type", "paragraph")  # heading, paragraph, list
                })
                content_parts.append(content)

        return {
            "id": document.id,
            "filename": document.filename,
            "content_type": document.content_type,
            "file_size": document.file_size,
            "chunk_count": document.chunk_count,
            "full_content": "\n\n".join(content_parts),
            "chunks": chunk_contents
        }
    
    async def get_document_chunks(self, document_id: str) -> List[Dict[str, Any]]:
        """
        Get document chunks for preview.

        Args:
            document_id: ID of the document

        Returns:
            List of chunk metadata (id, chunk_index, filename)

        Raises:
            ValueError: If document not found
        """
        chunks = await self._fetch_chunks_from_pinecone(document_id)
        if not chunks:
            raise ValueError(f"Document not found: {document_id}")

        # Extract filename from first chunk
        default_filename = chunks[0].get("metadata", {}).get("filename", "Unknown Document")

        # Build response (chunks already sorted by _fetch_chunks_from_pinecone)
        return [
            {
                "id": chunk["id"],
                "chunk_index": chunk.get("metadata", {}).get("chunk_index", 0),
                "filename": chunk.get("metadata", {}).get("filename", default_filename),
                "content_type": chunk.get("metadata", {}).get("content_type", "paragraph")
            }
            for chunk in chunks
            if chunk.get("metadata")
        ]


# Singleton instance
_document_service: Optional[DocumentService] = None

def get_document_service() -> DocumentService:
    """Get the singleton DocumentService instance."""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
